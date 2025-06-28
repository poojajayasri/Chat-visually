# src/services/document_service.py
import hashlib
import mimetypes
from abc import ABC, abstractmethod
from dataclasses import dataclass
from pathlib import Path
from typing import List, Optional, Dict, Any, BinaryIO
import uuid
from datetime import datetime

import streamlit as st
from pypdf import PdfReader
from docx import Document as DocxDocument
import pandas as pd
from PIL import Image
import pytesseract
from youtube_transcript_api import YouTubeTranscriptApi
import requests
from bs4 import BeautifulSoup

from ..config import StorageConfig
from ..models.document import Document, DocumentChunk, DocumentMetadata
from ..utils.logger import get_logger
from ..utils.text_splitter import SmartTextSplitter

logger = get_logger(__name__)

@dataclass
class ProcessingResult:
    """Result of document processing."""
    success: bool
    document: Optional[Document] = None
    error: Optional[str] = None
    chunks_created: int = 0

class DocumentProcessor(ABC):
    """Abstract base class for document processors."""
    
    @abstractmethod
    def can_process(self, file_path: Path, mime_type: str) -> bool:
        """Check if this processor can handle the file type."""
        pass
    
    @abstractmethod
    def process(self, file_path: Path, metadata: DocumentMetadata) -> List[DocumentChunk]:
        """Process the document and return chunks."""
        pass

class PDFProcessor(DocumentProcessor):
    """PDF document processor."""
    
    def can_process(self, file_path: Path, mime_type: str) -> bool:
        return file_path.suffix.lower() == '.pdf' or mime_type == 'application/pdf'
    
    def process(self, file_path: Path, metadata: DocumentMetadata) -> List[DocumentChunk]:
        """Process PDF document."""
        try:
            reader = PdfReader(file_path)
            text_parts = []
            
            for page_num, page in enumerate(reader.pages, 1):
                text = page.extract_text()
                if text.strip():
                    text_parts.append({
                        'content': text.strip(),
                        'page': page_num,
                        'type': 'text'
                    })
            
            # Split into chunks
            splitter = SmartTextSplitter(
                chunk_size=1000,
                chunk_overlap=200,
                separators=["\n\n", "\n", ". ", "! ", "? "]
            )
            
            chunks = []
            for part in text_parts:
                text_chunks = splitter.split_text(part['content'])
                for i, chunk_text in enumerate(text_chunks):
                    chunk = DocumentChunk(
                        id=str(uuid.uuid4()),
                        document_id=metadata.id,
                        content=chunk_text,
                        metadata={
                            'page': part['page'],
                            'chunk_index': i,
                            'source_type': 'pdf'
                        }
                    )
                    chunks.append(chunk)
            
            return chunks
            
        except Exception as e:
            logger.error(f"Error processing PDF {file_path}: {e}")
            raise

class DocxProcessor(DocumentProcessor):
    """DOCX document processor."""
    
    def can_process(self, file_path: Path, mime_type: str) -> bool:
        return (file_path.suffix.lower() == '.docx' or 
                mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    
    def process(self, file_path: Path, metadata: DocumentMetadata) -> List[DocumentChunk]:
        """Process DOCX document."""
        try:
            doc = DocxDocument(file_path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            text_content = '\n'.join(paragraphs)
            
            splitter = SmartTextSplitter(chunk_size=1000, chunk_overlap=200)
            text_chunks = splitter.split_text(text_content)
            
            chunks = []
            for i, chunk_text in enumerate(text_chunks):
                chunk = DocumentChunk(
                    id=str(uuid.uuid4()),
                    document_id=metadata.id,
                    content=chunk_text,
                    metadata={
                        'chunk_index': i,
                        'source_type': 'docx'
                    }
                )
                chunks.append(chunk)
            
            return chunks
            
        except Exception as e:
            logger.error(f"Error processing DOCX {file_path}: {e}")
            raise

class ImageProcessor(DocumentProcessor):
    """Image document processor using OCR."""
    
    def can_process(self, file_path: Path, mime_type: str) -> bool:
        return (file_path.suffix.lower() in ['.png', '.jpg', '.jpeg'] or
                mime_type.startswith('image/'))
    
    def process(self, file_path: Path, metadata: DocumentMetadata) -> List[DocumentChunk]:
        """Process image using OCR."""
        try:
            image = Image.open(file_path)
            text = pytesseract.image_to_string(image)
            
            if not text.strip():
                return []
            
            splitter = SmartTextSplitter(chunk_size=1000, chunk_overlap=200)
            text_chunks = splitter.split_text(text)
            
            chunks = []
            for i, chunk_text in enumerate(text_chunks):
                chunk = DocumentChunk(
                    id=str(uuid.uuid4()),
                    document_id=metadata.id,
                    content=chunk_text,
                    metadata={
                        'chunk_index': i,
                        'source_type': 'image_ocr'
                    }
                )
                chunks.append(chunk)
            
            return chunks
            
        except Exception as e:
            logger.error(f"Error processing image {file_path}: {e}")
            raise

class CSVProcessor(DocumentProcessor):
    """CSV document processor."""
    
    def can_process(self, file_path: Path, mime_type: str) -> bool:
        return file_path.suffix.lower() == '.csv' or mime_type == 'text/csv'
    
    def process(self, file_path: Path, metadata: DocumentMetadata) -> List[DocumentChunk]:
        """Process CSV document."""
        try:
            df = pd.read_csv(file_path)
            
            # Convert DataFrame to text representation
            text_content = f"CSV Data Summary:\n"
            text_content += f"Columns: {', '.join(df.columns)}\n"
            text_content += f"Rows: {len(df)}\n\n"
            
            # Add column descriptions
            for col in df.columns:
                text_content += f"Column '{col}':\n"
                text_content += f"- Data type: {df[col].dtype}\n"
                text_content += f"- Non-null values: {df[col].count()}\n"
                if df[col].dtype in ['int64', 'float64']:
                    text_content += f"- Mean: {df[col].mean():.2f}\n"
                text_content += f"- Sample values: {df[col].head(3).tolist()}\n\n"
            
            # Add first few rows as text
            text_content += "Sample Data:\n"
            text_content += df.head(10).to_string()
            
            splitter = SmartTextSplitter(chunk_size=1500, chunk_overlap=300)
            text_chunks = splitter.split_text(text_content)
            
            chunks = []
            for i, chunk_text in enumerate(text_chunks):
                chunk = DocumentChunk(
                    id=str(uuid.uuid4()),
                    document_id=metadata.id,
                    content=chunk_text,
                    metadata={
                        'chunk_index': i,
                        'source_type': 'csv',
                        'rows': len(df),
                        'columns': list(df.columns)
                    }
                )
                chunks.append(chunk)
            
            return chunks
            
        except Exception as e:
            logger.error(f"Error processing CSV {file_path}: {e}")
            raise

class DocumentService:
    """Service for managing document processing and storage."""
    
    def __init__(self, config: StorageConfig):
        self.config = config
        self.processors = [
            PDFProcessor(),
            DocxProcessor(),
            ImageProcessor(),
            CSVProcessor(),
        ]
    
    def validate_file(self, uploaded_file) -> tuple[bool, Optional[str]]:
        """Validate uploaded file."""
        if uploaded_file.size > self.config.max_file_size:
            return False, f"File size exceeds {self.config.max_file_size / (1024*1024):.1f}MB limit"
        
        file_extension = Path(uploaded_file.name).suffix.lower()
        if file_extension not in self.config.allowed_extensions:
            return False, f"File type {file_extension} not supported"
        
        return True, None
    
    def save_uploaded_file(self, uploaded_file, user_id: str) -> Path:
        """Save uploaded file to storage."""
        # Create user-specific directory
        user_dir = self.config.upload_dir / user_id
        user_dir.mkdir(exist_ok=True)
        
        # Generate unique filename
        file_hash = hashlib.md5(uploaded_file.read()).hexdigest()[:8]
        uploaded_file.seek(0)  # Reset file pointer
        
        filename = f"{file_hash}_{uploaded_file.name}"
        file_path = user_dir / filename
        
        # Save file
        with open(file_path, "wb") as f:
            f.write(uploaded_file.read())
        
        return file_path
    
    def process_document(self, file_path: Path, user_id: str, 
                        title: Optional[str] = None) -> ProcessingResult:
        """Process a document and return chunks."""
        try:
            # Determine MIME type
            mime_type, _ = mimetypes.guess_type(file_path)
            
            # Find appropriate processor
            processor = None
            for p in self.processors:
                if p.can_process(file_path, mime_type or ""):
                    processor = p
                    break
            
            if not processor:
                return ProcessingResult(
                    success=False,
                    error=f"No processor found for file type: {file_path.suffix}"
                )
            
            # Create document metadata
            metadata = DocumentMetadata(
                id=str(uuid.uuid4()),
                title=title or file_path.stem,
                filename=file_path.name,
                file_type=file_path.suffix,
                mime_type=mime_type,
                file_size=file_path.stat().st_size,
                user_id=user_id,
                created_at=datetime.now(),
                file_path=str(file_path)
            )
            
            # Process document
            chunks = processor.process(file_path, metadata)
            
            # Create document object
            document = Document(
                metadata=metadata,
                chunks=chunks
            )
            
            return ProcessingResult(
                success=True,
                document=document,
                chunks_created=len(chunks)
            )
            
        except Exception as e:
            logger.error(f"Error processing document {file_path}: {e}")
            return ProcessingResult(
                success=False,
                error=str(e)
            )
    
    def process_url(self, url: str, user_id: str, url_type: str = "web") -> ProcessingResult:
        """Process content from URL (YouTube or web)."""
        try:
            if "youtube.com" in url or "youtu.be" in url:
                return self._process_youtube_url(url, user_id)
            else:
                return self._process_web_url(url, user_id)
        except Exception as e:
            logger.error(f"Error processing URL {url}: {e}")
            return ProcessingResult(success=False, error=str(e))
    
    def _process_youtube_url(self, url: str, user_id: str) -> ProcessingResult:
        """Process YouTube video transcript."""
        try:
            # Extract video ID
            video_id = None
            if "youtube.com" in url:
                video_id = url.split("v=")[1].split("&")[0]
            elif "youtu.be" in url:
                video_id = url.split("/")[-1]
            
            if not video_id:
                return ProcessingResult(success=False, error="Invalid YouTube URL")
            
            # Get transcript
            transcript = YouTubeTranscriptApi.get_transcript(video_id)
            text_content = " ".join([entry['text'] for entry in transcript])
            
            # Create metadata
            metadata = DocumentMetadata(
                id=str(uuid.uuid4()),
                title=f"YouTube Video: {video_id}",
                filename=f"youtube_{video_id}.txt",
                file_type=".txt",
                mime_type="text/plain",
                file_size=len(text_content.encode()),
                user_id=user_id,
                created_at=datetime.now(),
                source_url=url
            )
            
            # Split into chunks
            splitter = SmartTextSplitter(chunk_size=1000, chunk_overlap=200)
            text_chunks = splitter.split_text(text_content)
            
            chunks = []
            for i, chunk_text in enumerate(text_chunks):
                chunk = DocumentChunk(
                    id=str(uuid.uuid4()),
                    document_id=metadata.id,
                    content=chunk_text,
                    metadata={
                        'chunk_index': i,
                        'source_type': 'youtube',
                        'video_id': video_id
                    }
                )
                chunks.append(chunk)
            
            document = Document(metadata=metadata, chunks=chunks)
            
            return ProcessingResult(
                success=True,
                document=document,
                chunks_created=len(chunks)
            )
            
        except Exception as e:
            logger.error(f"Error processing YouTube URL {url}: {e}")
            return ProcessingResult(success=False, error=str(e))
    
    def _process_web_url(self, url: str, user_id: str) -> ProcessingResult:
        """Process web page content."""
        try:
            # Fetch web content
            headers = {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
            }
            response = requests.get(url, headers=headers, timeout=10)
            response.raise_for_status()
            
            # Parse content
            soup = BeautifulSoup(response.content, 'html.parser')
            
            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()
            
            # Get text content
            text_content = soup.get_text()
            
            # Clean up text
            lines = (line.strip() for line in text_content.splitlines())
            chunks_text = (phrase.strip() for line in lines for phrase in line.split("  "))
            text_content = ' '.join(chunk for chunk in chunks_text if chunk)
            
            if not text_content.strip():
                return ProcessingResult(success=False, error="No text content found in URL")
            
            # Create metadata
            title = soup.find('title')
            title_text = title.get_text() if title else f"Web Page: {url}"
            
            metadata = DocumentMetadata(
                id=str(uuid.uuid4()),
                title=title_text,
                filename=f"webpage_{hashlib.md5(url.encode()).hexdigest()[:8]}.txt",
                file_type=".txt",
                mime_type="text/plain",
                file_size=len(text_content.encode()),
                user_id=user_id,
                created_at=datetime.now(),
                source_url=url
            )
            
            # Split into chunks
            splitter = SmartTextSplitter(chunk_size=1000, chunk_overlap=200)
            text_chunks = splitter.split_text(text_content)
            
            chunks = []
            for i, chunk_text in enumerate(text_chunks):
                chunk = DocumentChunk(
                    id=str(uuid.uuid4()),
                    document_id=metadata.id,
                    content=chunk_text,
                    metadata={
                        'chunk_index': i,
                        'source_type': 'web',
                        'url': url
                    }
                )
                chunks.append(chunk)
            
            document = Document(metadata=metadata, chunks=chunks)
            
            return ProcessingResult(
                success=True,
                document=document,
                chunks_created=len(chunks)
            )
            
        except Exception as e:
            logger.error(f"Error processing web URL {url}: {e}")
            return ProcessingResult(success=False, error=str(e))
